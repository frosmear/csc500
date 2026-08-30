from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import subprocess


class HomeworkPackager:

    def __init__(self, config_file="make_docx.cfg"):
        self.config_file = Path(config_file)
        self.config = self.load_config()

        self.assignment_title = self.config["assignment_title"]
        self.student_name = self.config["student_name"]

        self.pseudocode_file = Path(
            self.config["pseudocode_file"]
        )

        self.screenshot_directory = Path(
            self.config["screenshot_directory"]
        )

        self.output_file = Path(
            self.config["output_file"]
        )

        self.source_code_url = self.config["source_code_url"]

        self.document = Document()


    # ========================================================
    # CONFIGURATION
    # ========================================================

    def load_config(self):
        """Load settings from the configuration file."""

        config = {}

        with self.config_file.open(
            "r",
            encoding="utf-8"
        ) as file:

            for line in file:
                line = line.strip()

                if not line or line.startswith("#"):
                    continue

                if "=" not in line:
                    continue

                key, value = line.split("=", 1)

                config[key.strip()] = value.strip()

        return config


    # ========================================================
    # DOCUMENT CONTENT
    # ========================================================

    def add_title(self):
        """Add assignment title and student name."""

        self.document.add_heading(
            self.assignment_title,
            level=1
        )

        self.document.add_paragraph(
            f"Student: {self.student_name}"
        )


    def add_pseudocode(self):
        """Add pseudocode from the configured Markdown file."""

        self.document.add_heading(
            "Pseudocode",
            level=1
        )

        if not self.pseudocode_file.exists():
            self.document.add_paragraph(
                f"Pseudocode file not found: "
                f"{self.pseudocode_file}"
            )
            return

        with self.pseudocode_file.open(
            "r",
            encoding="utf-8"
        ) as file:

            for line in file:
                line = line.rstrip()

                if not line:
                    self.document.add_paragraph()
                    continue

                # Markdown headings
                if line.startswith("### "):
                    self.document.add_heading(
                        line[4:],
                        level=3
                    )

                elif line.startswith("## "):
                    self.document.add_heading(
                        line[3:],
                        level=2
                    )

                elif line.startswith("# "):
                    self.document.add_heading(
                        line[2:],
                        level=1
                    )

                # Bullet list
                elif line.startswith("- "):
                    self.document.add_paragraph(
                        line[2:],
                        style="List Bullet"
                    )

                # Numbered list
                elif (
                    len(line) > 2
                    and line[0].isdigit()
                    and ". " in line
                ):
                    self.document.add_paragraph(
                        line.split(". ", 1)[1],
                        style="List Number"
                    )

                # Normal paragraph
                else:
                    self.document.add_paragraph(line)


    def add_screenshots(self):
        """Add all PNG screenshots from the configured directory."""

        self.document.add_heading(
            "Screenshots",
            level=1
        )

        if not self.screenshot_directory.exists():
            self.document.add_paragraph(
                f"Screenshot directory not found: "
                f"{self.screenshot_directory}"
            )
            return

        screenshots = sorted(
            self.screenshot_directory.glob("*.png"),
            key=lambda path: path.name.lower()
        )

        if not screenshots:
            self.document.add_paragraph(
                "No PNG screenshots were found."
            )
            return

        for screenshot in screenshots:

            self.document.add_paragraph(
                screenshot.name
            )

            paragraph = self.document.add_paragraph()
            run = paragraph.add_run()

            run.add_picture(
                str(screenshot),
                width=Inches(6.0)
            )


    def add_source_code(self):
        """Add a hyperlink to the source code."""

        self.document.add_heading(
            "Source Code",
            level=1
        )

        paragraph = self.document.add_paragraph()

        self.add_hyperlink(
            paragraph,
            "View Source Code",
            self.source_code_url
        )


    # ========================================================
    # WORD / DOCX FUNCTIONS
    # ========================================================

    def add_hyperlink(self, paragraph, text, url):
        """Add a clickable hyperlink to a paragraph."""

        relationship_id = paragraph.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True
        )

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(
            qn("r:id"),
            relationship_id
        )

        run = OxmlElement("w:r")
        run_properties = OxmlElement("w:rPr")

        color = OxmlElement("w:color")
        color.set(
            qn("w:val"),
            "0563C1"
        )

        run_properties.append(color)

        underline = OxmlElement("w:u")
        underline.set(
            qn("w:val"),
            "single"
        )

        run_properties.append(underline)

        run.append(run_properties)

        text_element = OxmlElement("w:t")
        text_element.text = text

        run.append(text_element)
        hyperlink.append(run)

        paragraph._p.append(hyperlink)


    def save(self):
        """Save the completed Word document."""
        self.document.save(self.output_file)
        print(f"Created: {self.output_file.resolve()}"        )


    # ========================================================
    # GIT
    # ========================================================

    def git_commit(self):
        """Add the generated document and commit it."""

        if self.config.get("git_commit", "false").lower() != "true":
            print("Git commit deactivated in configuration.")
            return

        commit_message = self.config.get(
            "git_commit_message",
            f"Package {self.assignment_title}"
        )

        try:
            subprocess.run(
                ["git", "add", str(self.output_file)],
                check=True
            )

            subprocess.run(
                ["git", "commit", "-m", commit_message],
                check=True
            )

            print(f"Git commit created: {commit_message}")

        except subprocess.CalledProcessError as error:
            print(f"Git operation failed: {error}")


    # ========================================================
    # BUILD EVERYTHING
    # ========================================================

    def build(self):
        # Build the complete homework package.
        self.add_title()
        self.add_pseudocode()
        self.add_screenshots()
        self.add_source_code()
        self.save()
        self.git_commit()

# ============================================================
# MAIN
# ============================================================

def main():
    package = HomeworkPackager()
    package.build()

if __name__ == "__main__":
    main()
